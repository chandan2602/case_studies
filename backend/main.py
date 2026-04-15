import json
import os
import re
import tempfile
from collections import Counter
from datetime import date

import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
import numpy as np
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fastapi import FastAPI
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse

from sklearn.cluster import KMeans
from sklearn.ensemble import IsolationForest, RandomForestClassifier
from sklearn.metrics.pairwise import cosine_similarity
from sklearn.preprocessing import StandardScaler, LabelEncoder

app = FastAPI(title="RSB Case Study Generator")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["http://localhost:4200"],
    allow_methods=["*"],
    allow_headers=["*"],
)

DATA_PATH = os.path.join(os.path.dirname(__file__), "data", "newjsondata8 1.txt")

# Register reportlab for PDF if available
try:
    from reportlab.pdfbase import pdfmetrics
    _PDF_AVAILABLE = True
except Exception:
    _PDF_AVAILABLE = False

# ---------------------------------------------------------------------------
# Text sanitiser
# ---------------------------------------------------------------------------

def _clean(value) -> str:
    text = str(value) if value is not None else ""
    text = re.sub(r"[\u0900-\u097F]+", "", text)   # strip Devanagari
    text = re.sub(r"[^\x00-\x7F]+", "", text)       # strip other non-ASCII
    text = re.sub(r"\(\s*\)", "", text)              # remove empty parentheses ()
    text = re.sub(r"\s+", " ", text).strip()
    return text or "N/A"

FEATURE_COLS = [
    "total_beneficiaries_at_present", "total_beneficiaries_till_date",
    "coverage_noof_families", "coverage_noof_villages",
]

_df_cache = None

def load_data() -> pd.DataFrame:
    global _df_cache
    if _df_cache is not None:
        return _df_cache

    with open(DATA_PATH, encoding="utf-8") as f:
        records = json.load(f)
    df = pd.DataFrame(records)

    # ── Normalise column names to match what the pipeline expects ─────────────
    rename = {
        "surveycode":    "survey_code",
        "vertical":      "vertical_name",
        "subvertical":   "sub_vertical_name",
        "matrusansthan": "matrusansthan_name",
        "status":        "Status",
    }
    df.rename(columns=rename, inplace=True)

    # state_name already exists in JSON; use it as primary
    df = df[df["state_name"].notna()]

    # Add missing columns the pipeline references (set to 0 / empty)
    for col in ["total_volunteer_male", "total_volunteer_female",
                "total_salaried_male",  "total_salaried_female",
                "frequency_name", "geography_type_name",
                "building_type_name", "majority_population_comes_from"]:
        if col not in df.columns:
            df[col] = 0 if col not in ("frequency_name", "geography_type_name",
                                        "building_type_name", "majority_population_comes_from") else ""

    # Sanitise all string columns
    str_cols = df.select_dtypes(include="object").columns
    for col in str_cols:
        df[col] = df[col].apply(lambda v: _clean(v) if pd.notna(v) else v)

    # Numeric columns
    num_cols = [
        "total_beneficiaries_at_present", "total_beneficiaries_till_date",
        "beneficiaries_male_at_present",  "beneficiaries_female_at_present",
        "beneficiaries_boys_at_present",  "beneficiaries_girls_at_present",
        "beneficiaries_male_till_date",   "beneficiaries_female_till_date",
        "beneficiaries_boys_till_date",   "beneficiaries_girls_till_date",
        "coverage_noof_villages", "coverage_noof_families",
        "gram_samiti_karyakarta_male", "gram_samiti_karyakarta_female",
        "total_volunteer_male", "total_volunteer_female",
        "total_salaried_male",  "total_salaried_female",
    ]
    for col in num_cols:
        if col in df.columns:
            df[col] = pd.to_numeric(df[col], errors="coerce").fillna(0)

    _df_cache = df
    return df


# ---------------------------------------------------------------------------
# Routes
# ---------------------------------------------------------------------------

@app.get("/get_states")
async def get_states():
    df = load_data()
    states = sorted(df["state_name"].unique().tolist())
    return JSONResponse({"states": states})


@app.get("/get_districts")
async def get_districts(state: str):
    df = load_data()
    districts = sorted(df[df["state_name"] == state]["district_name"].dropna().unique().tolist())
    return JSONResponse({"districts": districts})


@app.get("/get_villages")
async def get_villages(state: str, district: str):
    df = load_data()
    villages = sorted(
        df[(df["state_name"] == state) & (df["district_name"] == district)]["village_name"]
        .dropna().unique().tolist()
    )
    return JSONResponse({"villages": villages})


from pydantic import BaseModel

class ReportRequest(BaseModel):
    state: str
    district: str = ""
    village: str = ""
    format: str = "docx"

@app.post("/generate_case_study")
async def generate_case_study(body: ReportRequest):
    state    = body.state
    district = body.district
    village  = body.village
    format   = body.format
    df = load_data()

    # Filter based on what was selected
    mask = df["state_name"] == state
    if district:
        mask &= df["district_name"] == district
    if village:
        mask &= df["village_name"] == village
    filtered = df[mask].copy()

    if filtered.empty:
        return JSONResponse({"error": "No records found for the selected filters."}, status_code=404)

    scope = village or district or state
    ml = run_ml_pipeline(df, filtered, scope, state, district, village)
    charts = build_charts(filtered, ml, scope)

    # Fetch photos from set15 for matching survey codes
    survey_codes = filtered["survey_code"].dropna().unique().tolist() if "survey_code" in filtered.columns else []
    photos = fetch_photos_for_scope(survey_codes, max_photos=6)

    fmt = format.lower().strip()
    if fmt == "pdf":
        if not _PDF_AVAILABLE:
            return JSONResponse({"error": "PDF export requires reportlab. Run: pip install reportlab"}, status_code=500)
        out_path = build_pdf(state, district, village, scope, filtered, ml, charts, photos)
        media_type = "application/pdf"
        ext = "pdf"
    else:
        out_path = build_document(state, district, village, scope, filtered, ml, charts, photos)
        media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        ext = "docx"

    for p in charts.values():
        if os.path.exists(p):
            os.unlink(p)
    for _, _, tmp_path in photos:
        if os.path.exists(tmp_path):
            os.unlink(tmp_path)

    safe = re.sub(r"[^\w]", "_", scope)
    return FileResponse(
        out_path,
        media_type=media_type,
        filename=f"case_study_{safe}.{ext}",
    )


# ---------------------------------------------------------------------------
# ML + Analysis Pipeline
# ---------------------------------------------------------------------------

def run_ml_pipeline(df, filtered, scope, state, district, village) -> dict:
    r = {}
    scaler = StandardScaler()

    # ── Scope metadata ────────────────────────────────────────────────────────
    r["scope_level"] = "Village" if village else ("District" if district else "State")
    r["total_records"] = len(filtered)
    r["verified"] = int((filtered["Status"] == "Verified").sum())
    r["not_verified"] = int((filtered["Status"] == "Not Verified").sum())
    r["verification_rate"] = round(r["verified"] / r["total_records"] * 100, 1)

    # ── Districts / Villages covered ──────────────────────────────────────────
    r["districts_covered"] = filtered["district_name"].nunique()
    r["villages_covered_unique"] = filtered["village_name"].nunique()

    # ── Vertical & sub-vertical breakdown ────────────────────────────────────
    r["verticals"] = filtered["vertical_name"].dropna().value_counts().to_dict()
    r["top_vertical"] = filtered["vertical_name"].dropna().value_counts().idxmax() \
        if filtered["vertical_name"].dropna().any() else "N/A"
    r["sub_verticals"] = filtered["sub_vertical_name"].dropna().value_counts().head(8).to_dict()

    # ── Beneficiary numbers ───────────────────────────────────────────────────
    r["ben_male_now"]   = int(filtered["beneficiaries_male_at_present"].sum())
    r["ben_female_now"] = int(filtered["beneficiaries_female_at_present"].sum())
    r["ben_boys_now"]   = int(filtered["beneficiaries_boys_at_present"].sum())
    r["ben_girls_now"]  = int(filtered["beneficiaries_girls_at_present"].sum())
    r["ben_total_now"]  = int(filtered["total_beneficiaries_at_present"].sum())

    r["ben_male_ever"]   = int(filtered["beneficiaries_male_till_date"].sum())
    r["ben_female_ever"] = int(filtered["beneficiaries_female_till_date"].sum())
    r["ben_boys_ever"]   = int(filtered["beneficiaries_boys_till_date"].sum())
    r["ben_girls_ever"]  = int(filtered["beneficiaries_girls_till_date"].sum())
    r["ben_total_ever"]  = int(filtered["total_beneficiaries_till_date"].sum())

    # Gender ratio
    total_now = r["ben_total_now"] or 1
    r["female_pct"] = round((r["ben_female_now"] + r["ben_girls_now"]) / total_now * 100, 1)
    r["male_pct"]   = round(100 - r["female_pct"], 1)

    # ── Human resources ───────────────────────────────────────────────────────
    r["vol_male"]    = int(filtered["total_volunteer_male"].sum())
    r["vol_female"]  = int(filtered["total_volunteer_female"].sum())
    r["vol_total"]   = r["vol_male"] + r["vol_female"]
    r["sal_male"]    = int(filtered["total_salaried_male"].sum())
    r["sal_female"]  = int(filtered["total_salaried_female"].sum())
    r["sal_total"]   = r["sal_male"] + r["sal_female"]
    r["gram_male"]   = int(filtered["gram_samiti_karyakarta_male"].sum())
    r["gram_female"] = int(filtered["gram_samiti_karyakarta_female"].sum())
    # Incharge count from JSON (one per project record)
    r["incharge_count"] = int(filtered["incharge_name"].notna().sum()) if "incharge_name" in filtered.columns else 0

    # ── Coverage ──────────────────────────────────────────────────────────────
    r["families_covered"] = int(filtered["coverage_noof_families"].sum())
    r["villages_covered"] = int(filtered["coverage_noof_villages"].sum())

    # ── Operational profile ───────────────────────────────────────────────────
    r["frequency_dist"]  = filtered["frequency_name"].dropna().value_counts().to_dict()
    r["geography_dist"]  = filtered["geography_type_name"].dropna().value_counts().to_dict()
    r["building_dist"]   = filtered["building_type_name"].dropna().value_counts().to_dict()
    r["population_dist"] = filtered["majority_population_comes_from"].dropna().value_counts().to_dict()
    r["surveytype_dist"] = filtered["surveytype_name"].dropna().value_counts().to_dict()

    # ── Programme maturity ────────────────────────────────────────────────────
    est = pd.to_numeric(filtered["established_in"], errors="coerce").dropna()
    if not est.empty:
        r["oldest_year"]  = int(est.min())
        r["newest_year"]  = int(est.max())
        r["avg_age_yrs"]  = round(date.today().year - est.mean(), 1)
    else:
        r["oldest_year"] = r["newest_year"] = r["avg_age_yrs"] = "N/A"

    # ── Matrusansthan (parent org) diversity ──────────────────────────────────
    r["matrusansthan_count"] = filtered["matrusansthan_name"].nunique()
    r["top_matrusansthan"]   = filtered["matrusansthan_name"].dropna().value_counts().to_dict()

    # ── Text column analysis (activities, impact, challenges, achievements) ───
    r["text_analysis"] = _analyze_text_columns(filtered)

    # ── 1. Isolation Forest ───────────────────────────────────────────────────
    if len(filtered) >= 5:
        scaled_f = scaler.fit_transform(filtered[FEATURE_COLS])
        iso = IsolationForest(contamination=0.1, random_state=42)
        filtered["anomaly"] = iso.fit_predict(scaled_f)
        r["outlier_count"] = int((filtered["anomaly"] == -1).sum())
        clean = filtered[filtered["anomaly"] == 1].copy()
    else:
        r["outlier_count"] = 0
        clean = filtered.copy()
    r["clean_count"] = len(clean)

    # ── 2. KMeans — segment projects ─────────────────────────────────────────
    n_clusters = min(3, len(clean))
    r["segments"] = []
    if len(clean) >= max(n_clusters, 2):
        scaled_c = scaler.fit_transform(clean[FEATURE_COLS])
        kmeans = KMeans(n_clusters=n_clusters, random_state=42, n_init=10)
        clean["cluster"] = kmeans.fit_predict(scaled_c)
        filtered.loc[clean.index, "cluster"] = clean["cluster"]
        seg = clean.groupby("cluster").agg(
            count=("total_beneficiaries_at_present", "count"),
            avg_ben=("total_beneficiaries_at_present", "mean"),
            avg_vol=("total_beneficiaries_till_date", "mean"),
            total_reach=("total_beneficiaries_till_date", "sum"),
            top_vertical=("vertical_name", lambda x: x.dropna().value_counts().idxmax()
                          if not x.dropna().empty else "N/A"),
        ).reset_index().sort_values("avg_ben")
        seg["label"] = ["Small Scale", "Medium Scale", "Large Scale"][:n_clusters]
        r["segments"] = seg.to_dict("records")

    # ── 3. Impact Scoring ─────────────────────────────────────────────────────
    for col in ["total_beneficiaries_till_date", "total_beneficiaries_at_present", "coverage_noof_families"]:
        mn, mx = filtered[col].min(), filtered[col].max()
        filtered[f"{col}_score"] = (filtered[col] - mn) / (mx - mn) if mx != mn else 0.5
    filtered["impact_score"] = (
        0.5 * filtered["total_beneficiaries_till_date_score"] +
        0.3 * filtered["total_beneficiaries_at_present_score"] +
        0.2 * filtered["coverage_noof_families_score"]
    )
    r["avg_impact_score"]   = round(filtered["impact_score"].mean(), 3)
    # Dynamic threshold: top 25% of scores in this dataset
    threshold = float(filtered["impact_score"].quantile(0.75))
    r["high_impact_count"]  = int((filtered["impact_score"] >= threshold).sum())
    r["top_projects"] = (
        filtered.nlargest(5, "impact_score")[
            ["prakalp_name", "vertical_name", "district_name",
             "total_beneficiaries_till_date", "total_beneficiaries_at_present", "impact_score"]
        ].fillna("N/A").to_dict("records")
    )

    # ── 4. Random Forest — predict vertical ───────────────────────────────────
    train_df = df[df["vertical_name"].notna()].copy()
    for c in FEATURE_COLS:
        train_df[c] = pd.to_numeric(train_df[c], errors="coerce").fillna(0)
    if len(train_df) >= 20:
        le = LabelEncoder()
        X = train_df[FEATURE_COLS].values
        y = le.fit_transform(train_df["vertical_name"])
        rf = RandomForestClassifier(n_estimators=100, random_state=42, n_jobs=-1)
        rf.fit(X, y)
        X_pred = filtered[FEATURE_COLS].values
        preds = le.inverse_transform(rf.predict(X_pred))
        filtered["predicted_vertical"] = preds
        pred_counts = pd.Series(preds).value_counts()
        r["predicted_top_vertical"] = pred_counts.idxmax()
        r["predicted_dist"] = pred_counts.to_dict()
        r["feature_importance"] = dict(zip(FEATURE_COLS, rf.feature_importances_.round(3)))
        match = (filtered["predicted_vertical"] == filtered["vertical_name"]).mean()
        r["prediction_match_rate"] = round(match * 100, 1)
    else:
        r["predicted_top_vertical"] = None
        r["feature_importance"] = {}
        r["prediction_match_rate"] = None

    # ── 5. Cosine Similarity — similar districts/states ───────────────────────
    group_col = "district_name" if r["scope_level"] == "State" else "village_name"
    vp = df.groupby(group_col).agg(
        total_projects=("survey_code", "count"),
        avg_ben=("total_beneficiaries_at_present", "mean"),
        avg_vol=("total_beneficiaries_till_date", "mean"),
        total_reach=("total_beneficiaries_till_date", "sum"),
    ).reset_index().fillna(0)
    sim_features = ["total_projects", "avg_ben", "avg_vol", "total_reach"]
    sim_matrix = StandardScaler().fit_transform(vp[sim_features])
    target_val = district if r["scope_level"] == "State" else village
    target_idx = vp[vp[group_col] == target_val].index if target_val else []
    if len(target_idx) > 0:
        sims = cosine_similarity([sim_matrix[target_idx[0]]], sim_matrix)[0]
        vp["similarity"] = sims
        similar = vp[vp[group_col] != target_val].sort_values("similarity", ascending=False).head(3)
        r["similar_units"] = similar[[group_col, "total_projects", "avg_ben", "similarity"]].to_dict("records")
        r["similar_col"] = group_col
    else:
        r["similar_units"] = []
        r["similar_col"] = group_col

    return r


def _analyze_text_columns(filtered: pd.DataFrame) -> dict:
    """
    Analyze ALL records and build aggregate summaries per text column.
    Uses real numbers from the dataset + best representative sentences.
    """
    result = {}
    text_cols = {
        "activities": "Key Activities",
        "impact": "Reported Impact",
        "achievements": "Achievements",
        "challenges": "Challenges",
        "reasons_for_opening_the_prakalp": "Reasons for Opening",
    }
    keywords = [
        "education", "health", "water", "women", "children", "youth",
        "skill", "livelihood", "awareness", "training", "community",
        "tribal", "rural", "nutrition", "sanitation", "employment",
        "cultural", "sports", "environment", "digital", "medical",
        "discipline", "confidence", "value", "support", "development",
    ]

    total_prakalps = len(filtered)
    total_ben_ever = int(filtered["total_beneficiaries_till_date"].sum()) if "total_beneficiaries_till_date" in filtered.columns else 0
    total_ben_now  = int(filtered["total_beneficiaries_at_present"].sum()) if "total_beneficiaries_at_present" in filtered.columns else 0
    centre_word    = "centres" if total_prakalps > 1 else "centre"

    for col, label in text_cols.items():
        if col not in filtered.columns:
            continue
        texts = filtered[col].dropna().astype(str)
        if texts.empty:
            continue

        filled = int(filtered[col].notna().sum())

        # Keyword frequency across ALL records
        lower_all = texts.str.lower()
        kw_counts = {}
        for kw in keywords:
            cnt = int(lower_all.str.contains(kw, na=False).sum())
            if cnt > 0:
                kw_counts[kw] = cnt
        top_kw = sorted(kw_counts.items(), key=lambda x: x[1], reverse=True)[:5]
        top_theme_words = [k for k, _ in top_kw[:3]]
        themes_str = ", ".join(top_theme_words) if top_theme_words else "community development"

        # Collect all sentences from all records
        all_sentences = []
        for raw in texts.tolist():
            cleaned = _clean(raw)
            if not cleaned or cleaned == "N/A":
                continue
            for sent in re.split(r"[.\n]+", cleaned):
                sent = sent.strip()
                if len(sent) > 40:
                    all_sentences.append(sent)

        # Score sentences by keyword hits + length
        def score_sentence(s):
            sl = s.lower()
            return sum(1 for kw in keywords if kw in sl) + min(len(s), 200) / 200.0

        scored = sorted(all_sentences, key=score_sentence, reverse=True)

        # Pick top 2 diverse sentences
        # If total beneficiaries is large, skip sentences mentioning small specific numbers
        small_number_threshold = total_ben_ever * 0.001  # less than 0.1% of total = "small"

        def _mentions_small_number(s):
            """Return True if sentence contains a standalone number smaller than threshold."""
            nums = re.findall(r'\b(\d+)\b', s)
            return any(int(n) < small_number_threshold for n in nums if int(n) > 0)

        selected = []
        for sent in scored:
            # Skip sentences with small numbers when dataset is large
            if total_ben_ever > 1000 and _mentions_small_number(sent):
                continue
            words = set(sent.lower().split())
            is_dup = any(
                len(words & set(s.lower().split())) / max(len(words), 1) > 0.6
                for s in selected
            )
            if not is_dup:
                selected.append(sent)
            if len(selected) == 2:
                break

        # Build aggregate summary using real numbers
        if col == "reasons_for_opening_the_prakalp":
            summary = (
                f"Across {filled} of {total_prakalps} prakalps, these {centre_word} were established "
                f"to provide access to {themes_str}. "
                f"Collectively they serve {total_ben_ever:,} cumulative beneficiaries."
            )
        elif col == "impact":
            summary = (
                f"Cumulatively, {total_ben_ever:,} beneficiaries have been reached across "
                f"{total_prakalps} prakalps ({total_ben_now:,} currently active). "
                f"Key areas of impact include {themes_str}. "
                f"These prakalps have created measurable change in the lives of individuals and communities."
            )
        elif col == "achievements":
            boys_ever  = int(filtered["beneficiaries_boys_till_date"].sum())  if "beneficiaries_boys_till_date"  in filtered.columns else 0
            girls_ever = int(filtered["beneficiaries_girls_till_date"].sum()) if "beneficiaries_girls_till_date" in filtered.columns else 0
            male_ever  = int(filtered["beneficiaries_male_till_date"].sum())  if "beneficiaries_male_till_date"  in filtered.columns else 0
            female_ever= int(filtered["beneficiaries_female_till_date"].sum())if "beneficiaries_female_till_date" in filtered.columns else 0
            summary = (
                f"Across all {total_prakalps} prakalps, a total of {total_ben_ever:,} individuals have benefited "
                f"({male_ever + boys_ever:,} male / {female_ever + girls_ever:,} female). "
                f"Notable achievements span {themes_str}, with {filled} prakalps reporting documented outcomes. "
                f"These prakalps have collectively transformed communities through sustained engagement, "
                f"skill development, and value-based education."
            )
        elif col == "challenges":
            summary = (
                f"{filled} of {total_prakalps} prakalps reported challenges. "
                f"Common themes include {themes_str}."
            )
        else:
            summary = (
                f"{filled} prakalps reported activities covering {themes_str}."
            )

        samples = [summary] + [
            s[:200].rsplit(" ", 1)[0] + "..." if len(s) > 200 else s
            for s in selected
        ]

        result[label] = {
            "filled_count": filled,
            "top_themes": top_kw,
            "samples": samples,
        }

    return result


_photo_df_cache = None


def load_photo_data() -> pd.DataFrame:
    global _photo_df_cache
    if _photo_df_cache is not None:
        return _photo_df_cache
    try:
        df = pd.read_csv(PHOTO_DATA_PATH, sep="\t")
        df.columns = [c.strip() for c in df.columns]
        # Build full URL = base URL + filename
        df["full_url"] = df["URL"].str.rstrip("/") + "/" + df["File Name"].str.strip()
        # Extract category from filename pattern: Photo_CODE_Category_timestamp.ext
        import re as _re
        def _cat(fname):
            m = _re.search(r"Photo_\w+_(.*?)_\d{8,}", str(fname))
            return m.group(1).strip() if m else ""
        df["category"] = df["File Name"].apply(_cat)
        _photo_df_cache = df
    except Exception as e:
        print(f"Could not load photo data: {e}")
        _photo_df_cache = pd.DataFrame(columns=["File Name", "Survey", "URL", "full_url", "category"])
    return _photo_df_cache


def fetch_photos_for_scope(survey_codes: list, max_photos: int = 6) -> list:
    """
    Returns list of (url, category, local_tmp_path) for photos matching survey codes.
    Downloads up to max_photos images, picking one per category where possible.
    """
    import urllib.request as _req
    photo_df = load_photo_data()
    if photo_df.empty or not survey_codes:
        return []

    subset = photo_df[photo_df["Survey"].isin(survey_codes)]
    if subset.empty:
        return []

    selected = []
    seen_cats = set()
    # First pass: one per priority category
    for cat in PHOTO_CATEGORIES:
        cat_rows = subset[subset["category"].str.contains(cat[:20], case=False, na=False)]
        if not cat_rows.empty:
            row = cat_rows.iloc[0]
            selected.append((row["full_url"], cat))
            seen_cats.add(cat)
        if len(selected) >= max_photos:
            break
    # Fill remaining slots from any category
    if len(selected) < max_photos:
        for _, row in subset.iterrows():
            if len(selected) >= max_photos:
                break
            if row["full_url"] not in [s[0] for s in selected]:
                selected.append((row["full_url"], row["category"]))

    # Download each photo to a temp file
    result = []
    for url, cat in selected:
        try:
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".jpg")
            _req.urlretrieve(url, tmp.name)
            result.append((url, _clean(cat) or "Project Photo", tmp.name))
        except Exception as e:
            print(f"Could not download photo {url}: {e}")
    return result


def build_charts(filtered, ml, scope) -> dict:
    charts = {}
    COLORS = ["#1f497d", "#2e75b6", "#4472c4", "#ed7d31", "#70ad47", "#ffc000", "#a9d18e"]

    def _save(fig):
        t = tempfile.NamedTemporaryFile(delete=False, suffix=".png")
        fig.savefig(t.name, dpi=150, bbox_inches="tight")
        plt.close(fig)
        return t.name

    # 1. Projects by Vertical — horizontal bar
    if ml["verticals"]:
        s = pd.Series(ml["verticals"]).sort_values()
        fig, ax = plt.subplots(figsize=(8, max(3, len(s) * 0.5)))
        bars = ax.barh(s.index, s.values, color=COLORS[:len(s)], edgecolor="white")
        ax.bar_label(bars, padding=3, fontsize=9)
        ax.set_title(f"Projects by Vertical — {_clean(scope)}", fontsize=12, fontweight="bold")
        ax.set_xlabel("No. of Projects")
        plt.tight_layout()
        charts["vertical_chart"] = _save(fig)

    # 2. Beneficiary gender pie (current + cumulative)
    male = ml["ben_male_now"] + ml["ben_boys_now"]
    female = ml["ben_female_now"] + ml["ben_girls_now"]
    if male + female > 0:
        male_e = ml["ben_male_ever"] + ml["ben_boys_ever"]
        female_e = ml["ben_female_ever"] + ml["ben_girls_ever"]
        fig, axes = plt.subplots(1, 2, figsize=(8, 4))
        for ax, vals, title in [
            (axes[0], [male, female],     "Current Beneficiaries"),
            (axes[1], [male_e, female_e], "Cumulative Beneficiaries"),
        ]:
            ax.pie(vals, labels=["Male", "Female"],
                   colors=["#1f497d", "#ed7d31"], autopct="%1.1f%%", startangle=90,
                   wedgeprops={"edgecolor": "white", "linewidth": 1.5})
            ax.set_title(title, fontsize=11, fontweight="bold")
        plt.suptitle("Gender Distribution", fontsize=12, fontweight="bold")
        plt.tight_layout()
        charts["gender_chart"] = _save(fig)

    # 3. Human resources — grouped bar (Male vs Female)
    hr = {
        "Volunteers":     (ml["vol_male"],  ml["vol_female"]),
        "Salaried Staff": (ml["sal_male"],  ml["sal_female"]),
        "Gram Samiti":    (ml["gram_male"], ml["gram_female"]),
    }
    hr = {k: v for k, v in hr.items() if sum(v) > 0}
    if hr:
        cats = list(hr.keys())
        males   = [hr[c][0] for c in cats]
        females = [hr[c][1] for c in cats]
        x = range(len(cats))
        w = 0.35
        fig, ax = plt.subplots(figsize=(7, 4))
        b1 = ax.bar([i - w/2 for i in x], males,   w, label="Male",   color="#1f497d", edgecolor="white")
        b2 = ax.bar([i + w/2 for i in x], females, w, label="Female", color="#ed7d31", edgecolor="white")
        ax.bar_label(b1, padding=2, fontsize=8)
        ax.bar_label(b2, padding=2, fontsize=8)
        ax.set_xticks(list(x)); ax.set_xticklabels(cats)
        ax.set_title("Human Resources — Male vs Female", fontsize=12, fontweight="bold")
        ax.set_ylabel("Count"); ax.legend()
        plt.tight_layout()
        charts["hr_chart"] = _save(fig)

    # 4. Verification status pie
    v, nv = ml["verified"], ml["not_verified"]
    if v + nv > 0:
        fig, ax = plt.subplots(figsize=(5, 4))
        ax.pie([v, nv], labels=["Verified", "Not Verified"],
               colors=["#70ad47", "#ed7d31"], autopct="%1.1f%%", startangle=90,
               wedgeprops={"edgecolor": "white", "linewidth": 1.5})
        ax.set_title("Project Verification Status", fontsize=12, fontweight="bold")
        plt.tight_layout()
        charts["verify_chart"] = _save(fig)

    # 5. Top sub-verticals horizontal bar
    if ml["sub_verticals"]:
        sv = pd.Series(ml["sub_verticals"]).sort_values(ascending=False).head(8)
        fig, ax = plt.subplots(figsize=(8, max(3, len(sv) * 0.55)))
        bars = ax.barh(sv.index[::-1], sv.values[::-1], color="#2e75b6", edgecolor="white")
        ax.bar_label(bars, padding=3, fontsize=9)
        ax.set_title("Top Sub-Verticals", fontsize=12, fontweight="bold")
        ax.set_xlabel("No. of Projects")
        plt.tight_layout()
        charts["subvertical_chart"] = _save(fig)

    # 6. Geography type pie
    if ml["geography_dist"] and len(ml["geography_dist"]) > 1:
        gd = pd.Series(ml["geography_dist"])
        fig, ax = plt.subplots(figsize=(5, 4))
        ax.pie(gd.values, labels=gd.index, colors=COLORS[:len(gd)],
               autopct="%1.1f%%", startangle=90,
               wedgeprops={"edgecolor": "white", "linewidth": 1.5})
        ax.set_title("Geography Type Distribution", fontsize=12, fontweight="bold")
        plt.tight_layout()
        charts["geo_chart"] = _save(fig)

    # 7. Beneficiary reach bar (current vs cumulative by category)
    ben_cats = ["Adult Male", "Adult Female", "Boys", "Girls"]
    ben_now  = [ml["ben_male_now"], ml["ben_female_now"], ml["ben_boys_now"], ml["ben_girls_now"]]
    ben_ever = [ml["ben_male_ever"], ml["ben_female_ever"], ml["ben_boys_ever"], ml["ben_girls_ever"]]
    if sum(ben_now) + sum(ben_ever) > 0:
        x = range(len(ben_cats))
        w = 0.35
        fig, ax = plt.subplots(figsize=(8, 4))
        b1 = ax.bar([i - w/2 for i in x], ben_now,  w, label="Current",    color="#1f497d", edgecolor="white")
        b2 = ax.bar([i + w/2 for i in x], ben_ever, w, label="Cumulative", color="#4472c4", edgecolor="white")
        ax.bar_label(b1, padding=2, fontsize=8, fmt="%d")
        ax.bar_label(b2, padding=2, fontsize=8, fmt="%d")
        ax.set_xticks(list(x)); ax.set_xticklabels(ben_cats)
        ax.set_title("Beneficiary Reach — Current vs Cumulative", fontsize=12, fontweight="bold")
        ax.set_ylabel("Beneficiaries"); ax.legend()
        plt.tight_layout()
        charts["reach_chart"] = _save(fig)

    return charts


# ---------------------------------------------------------------------------
# Document helpers
# ---------------------------------------------------------------------------

def _h(doc, text, level=1):
    p = doc.add_heading(_clean(text), level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

def _kv(doc, key, value):
    p = doc.add_paragraph()
    p.add_run(f"{_clean(key)}: ").bold = True
    p.add_run(_clean(value))

def _tbl(doc, headers, rows):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light List Accent 1"
    for i, h in enumerate(headers):
        t.rows[0].cells[i].text = _clean(h)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = _clean(v)

def _fig(doc, path, caption, width=5.5):
    p = doc.add_paragraph()
    p.add_run(f"Figure: {_clean(caption)}").italic = True
    doc.add_picture(path, width=Inches(width))


# ---------------------------------------------------------------------------
# Document builder
# ---------------------------------------------------------------------------

def build_document(state, district, village, scope, filtered, ml, charts, photos=None) -> str:
    from docx.oxml.ns import qn as _qn
    from docx.oxml import OxmlElement as _OxmlElement

    doc = Document()
    LOGO_PATH = os.path.join(os.path.dirname(__file__), "logo_dark.png")
    scope_line = _clean(village or (f"{district}, {state}" if district else state))

    # Compact margins
    for sec in doc.sections:
        sec.top_margin = sec.bottom_margin = Inches(0.6)
        sec.left_margin = sec.right_margin = Inches(0.75)

    BLUE = RGBColor(0x1F, 0x49, 0x7D)

    def _p(text="", bold=False, italic=False, size=9, color=None, center=False, sa=3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(sa)
        p.paragraph_format.space_before = Pt(0)
        if center: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if text:
            r = p.add_run(str(text))
            r.bold = bold; r.italic = italic; r.font.size = Pt(size)
            if color: r.font.color.rgb = color
        return p

    def _sec(label):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after  = Pt(2)
        r = p.add_run(label)
        r.bold = True; r.font.size = Pt(10); r.font.color.rgb = BLUE
        pPr = p._p.get_or_add_pPr()
        pBdr = _OxmlElement("w:pBdr")
        bot = _OxmlElement("w:bottom")
        bot.set(_qn("w:val"), "single"); bot.set(_qn("w:sz"), "4")
        bot.set(_qn("w:space"), "1"); bot.set(_qn("w:color"), "1F497D")
        pBdr.append(bot); pPr.append(pBdr)

    def _kvrow(pairs, size=8.5):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        for i, (k, v) in enumerate(pairs):
            if i: p.add_run("   ").font.size = Pt(size)
            kr = p.add_run(f"{k}: "); kr.bold = True; kr.font.size = Pt(size)
            vr = p.add_run(str(v)); vr.font.size = Pt(size)

    def _stbl(headers, rows, fs=8):
        t = doc.add_table(rows=1, cols=len(headers))
        t.style = "Light List Accent 1"; t.autofit = True
        for i, h in enumerate(headers):
            t.rows[0].cells[i].text = ""
            r = t.rows[0].cells[i].paragraphs[0].add_run(_clean(h))
            r.bold = True; r.font.size = Pt(fs)
        for row in rows:
            cells = t.add_row().cells
            for i, v in enumerate(row):
                cells[i].text = ""
                cells[i].paragraphs[0].add_run(_clean(v)).font.size = Pt(fs)
        doc.add_paragraph().paragraph_format.space_after = Pt(2)

    def _img2col(key1, key2, w=2.9):
        t = doc.add_table(rows=1, cols=2); t.autofit = True
        for i, key in enumerate([key1, key2]):
            if key in charts:
                cell = t.rows[0].cells[i]
                cell.text = ""
                cp = cell.paragraphs[0]
                cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                cp.add_run().add_picture(charts[key], width=Inches(w))

    # ── PAGE 1: Cover ─────────────────────────────────────────────────────────
    if os.path.exists(LOGO_PATH):
        lp = doc.add_paragraph()
        lp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        lp.paragraph_format.space_before = Pt(50)
        lp.paragraph_format.space_after  = Pt(10)
        lp.add_run().add_picture(LOGO_PATH, width=Inches(2.0))
    _p("Sewa Sanskriti App", bold=True, size=18, color=BLUE, center=True, sa=6)
    _p("Case Study Report",  bold=True, size=22, color=BLUE, center=True, sa=4)
    _p(scope_line,           bold=True, size=14, color=BLUE, center=True, sa=4)
    # All matrusansthan names on cover
    all_mat = ", ".join(ml["top_matrusansthan"].keys())
    _p(f"Matrusansthans: {all_mat}", italic=True, size=8.5, center=True, sa=6)
    _p(f"Rashtriya Sewa Bharti (RSB)  |  Scope: {ml['scope_level']} — {scope_line}  |  "
       f"Date: {date.today().strftime('%B %d, %Y')}",
       italic=True, size=9, center=True, sa=0)
    doc.add_page_break()

    # ── PAGE 2: Overview + qualitative summary + 2 charts ────────────────────
    _sec("Overview")
    _kvrow([("Prakalps", f"{ml['total_records']:,}"),
            ("Verified", f"{ml['verified']:,} ({ml['verification_rate']}%)"),
            ("Districts", ml["districts_covered"]),
            ("Villages", ml["villages_covered_unique"])])
    _kvrow([("Cumulative Beneficiaries", f"{ml['ben_total_ever']:,}"),
            ("Incharge/Workers", f"{ml.get('incharge_count', ml['vol_total']):,}"),
            ("Gram Samiti", f"{ml['gram_male'] + ml['gram_female']:,}"),
            ("Families", f"{ml['families_covered']:,}")])
    mat_names = ", ".join(f"{k} ({v})" for k, v in ml["top_matrusansthan"].items())
    _kvrow([("Dominant Vertical", ml["top_vertical"]),
            ("Female Share", f"{ml['female_pct']}%"),
            ("Matrusansthans Involved", f"{ml['matrusansthan_count']} — {mat_names}")])

    # Qualitative highlights — Reasons, Impact, Achievements
    QUAL_SHOW = ["Reasons for Opening", "Reported Impact", "Achievements"]
    for label in QUAL_SHOW:
        data = ml["text_analysis"].get(label)
        if not data:
            continue
        _sec(label)
        if data.get("samples"):
            for s in data["samples"][:2]:
                p = doc.add_paragraph()
                p.paragraph_format.space_after  = Pt(2)
                p.paragraph_format.space_before = Pt(0)
                r = p.add_run(f"- {s}")
                r.font.size = Pt(8.5); r.italic = True
        if data["top_themes"]:
            themes_str = ", ".join(kw for kw, _ in data["top_themes"])
            _kvrow([("Key themes", themes_str)])

    _img2col("gender_chart", "vertical_chart", w=2.9)

    # ── PAGE 3: HR + Verticals + Operational ──────────────────────────────────
    _sec("Human Resources")
    hr_rows = []
    if ml["vol_total"] > 0:
        hr_rows.append(("Volunteers", str(ml["vol_male"]), str(ml["vol_female"]), str(ml["vol_total"])))
    if ml["sal_total"] > 0:
        hr_rows.append(("Salaried Staff", str(ml["sal_male"]), str(ml["sal_female"]), str(ml["sal_total"])))
    gram_total = ml["gram_male"] + ml["gram_female"]
    if gram_total > 0:
        hr_rows.append(("Gram Samiti", str(ml["gram_male"]), str(ml["gram_female"]), str(gram_total)))
    if ml.get("incharge_count", 0) > 0:
        hr_rows.append(("Incharge / Coordinators", str(ml["incharge_count"]), "-", str(ml["incharge_count"])))
    if hr_rows:
        _stbl(["Category", "Male", "Female", "Total"], hr_rows)
    else:
        _kvrow([("Note", "Worker breakdown not available in this dataset")])

    _sec("Beneficiary Reach")
    _stbl(["Category", "Current", "Cumulative"], [
        ("Male",   f"{ml['ben_male_now']:,}",   f"{ml['ben_male_ever']:,}"),
        ("Female", f"{ml['ben_female_now']:,}", f"{ml['ben_female_ever']:,}"),
        ("Boys",   f"{ml['ben_boys_now']:,}",   f"{ml['ben_boys_ever']:,}"),
        ("Girls",  f"{ml['ben_girls_now']:,}",  f"{ml['ben_girls_ever']:,}"),
        ("TOTAL",  f"{ml['ben_total_now']:,}",  f"{ml['ben_total_ever']:,}"),
    ])

    _sec("Service Verticals")
    _stbl(["Vertical", "Prakalps"],
          [(k, str(v)) for k, v in list(ml["verticals"].items())[:6]])

    # ── PAGE 4: Impact + Segments + Summary ───────────────────────────────────
    _sec("Top Impact Prakalps")
    _kvrow([("Avg Score", ml["avg_impact_score"]),
            ("High-Impact", ml["high_impact_count"]),
            ("Outliers", ml["outlier_count"])])
    if ml["top_projects"]:
        _stbl(["Prakalp Name", "Vertical", "Total Beneficiaries", "Score"],
              [(p["prakalp_name"][:28], p["vertical_name"],
                f"{int(p['total_beneficiaries_till_date']):,}",
                str(round(p["impact_score"], 2)))
               for p in ml["top_projects"][:5]])

    if ml["segments"]:
        _sec("Prakalp Segments")
        _stbl(["Segment", "Prakalps", "Avg Beneficiaries", "Top Vertical"],
              [(s["label"], str(s["count"]), str(round(s["avg_ben"], 0)), s["top_vertical"])
               for s in ml["segments"]])

    _sec("Recommendations")
    for i, rec in enumerate(_get_recommendations(ml)[:5], 1):
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.space_before = Pt(0)
        p.add_run(rec).font.size = Pt(8.5)

    _sec("Summary")
    scope_line_local = _clean(village or (f"{district}, {state}" if district else state))
    n = ml["total_records"]
    centre_word = "centres" if n > 1 else "centre"
    mat_names = ", ".join(ml["top_matrusansthan"].keys())
    verticals_list = ", ".join(list(ml["verticals"].keys())[:4])
    summary_text = (
        f"This {ml['scope_level'].lower()}-level report covers {n:,} prakalps operating across "
        f"{ml['districts_covered']} district(s) and {ml['villages_covered_unique']} village(s) in {scope_line_local}. "
        f"These {centre_word} are run under {ml['matrusansthan_count']} matrusansthans — {mat_names}. "
        f"Cumulatively, {ml['ben_total_ever']:,} beneficiaries have been served "
        f"({ml['ben_total_now']:,} currently active), with a female share of {ml['female_pct']}%. "
        f"The dominant service vertical is {ml['top_vertical']}, with prakalps also covering {verticals_list}. "
        f"{ml['high_impact_count']} prakalps were identified as high-impact based on beneficiary reach and coverage. "
        f"{ml['verified']:,} of {n:,} prakalps ({ml['verification_rate']}%) are verified."
    )
    for lbl in ["Reported Impact", "Achievements", "Reasons for Opening"]:
        data = ml["text_analysis"].get(lbl)
        if data and data.get("samples"):
            summary_text += f" {lbl}: {data['samples'][0]}"
            break
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.space_before = Pt(0)
    p.add_run(summary_text).font.size = Pt(8.5)

    # ── PAGE 5: Photos (2×2 grid, max 4) ─────────────────────────────────────
    if photos:
        doc.add_page_break()
        _sec("Photo Gallery")
        _p("Field photographs from project sites.", size=8, sa=4)
        for row_start in range(0, min(len(photos), 4), 2):
            pair = photos[row_start:row_start + 2]
            pt = doc.add_table(rows=1, cols=len(pair)); pt.autofit = True
            for i, (_, cap, tmp_path) in enumerate(pair):
                cell = pt.rows[0].cells[i]; cell.text = ""
                cp = cell.paragraphs[0]; cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
                try:
                    cp.add_run().add_picture(tmp_path, width=Inches(3.0))
                except Exception:
                    pass
                cap_p = cell.add_paragraph(_clean(cap))
                cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if cap_p.runs:
                    cap_p.runs[0].font.size = Pt(7); cap_p.runs[0].italic = True

    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".docx")
    doc.save(tmp.name)
    return tmp.name


# ---------------------------------------------------------------------------
# PDF builder (reportlab + Noto Sans Devanagari)
# ---------------------------------------------------------------------------

def build_pdf(state, district, village, scope, filtered, ml, charts, photos=None) -> str:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle,
        Image, PageBreak, HRFlowable,
    )
    from reportlab.lib.enums import TA_CENTER

    scope_line = _clean(village or (f"{district}, {state}" if district else state))
    LOGO_PATH  = os.path.join(os.path.dirname(__file__), "logo_dark.png")
    tmp = tempfile.NamedTemporaryFile(delete=False, suffix=".pdf")

    doc = SimpleDocTemplate(
        tmp.name, pagesize=A4,
        leftMargin=0.75*inch, rightMargin=0.75*inch,
        topMargin=0.6*inch,   bottomMargin=0.6*inch,
    )
    W    = A4[0] - 1.5*inch
    BLUE = colors.HexColor("#1f497d")

    def _s(name, **kw):
        kw.setdefault("fontName", "Helvetica")
        return ParagraphStyle(name, **kw)

    normal  = _s("pN",  fontSize=8.5, leading=12, spaceAfter=2)
    italic  = _s("pI",  fontSize=8.5, leading=12, spaceAfter=2, fontName="Helvetica-Oblique")
    kv_s    = _s("pKV", fontSize=8.5, leading=12, spaceAfter=2)
    sec_s   = _s("pSec",fontSize=10,  leading=13, spaceAfter=3, spaceBefore=5,
                 fontName="Helvetica-Bold", textColor=BLUE)
    title_s = _s("pTi", fontSize=22,  leading=28, spaceAfter=4,
                 fontName="Helvetica-Bold", textColor=BLUE, alignment=TA_CENTER)
    sub_s   = _s("pSu", fontSize=14,  leading=18, spaceAfter=4,
                 fontName="Helvetica-Bold", textColor=BLUE, alignment=TA_CENTER)
    meta_s  = _s("pMe", fontSize=9,   leading=13, spaceAfter=2, alignment=TA_CENTER)
    cap_s   = _s("pCa", fontSize=7,   leading=10, spaceAfter=2,
                 textColor=colors.grey, alignment=TA_CENTER)

    def _sec_elem(label):
        return [
            HRFlowable(width=W, thickness=0.5, color=BLUE, spaceAfter=2, spaceBefore=5),
            Paragraph(label, sec_s),
        ]

    def _kvrow_elem(pairs):
        parts = []
        for i, (k, v) in enumerate(pairs):
            if i:
                parts.append("   ")
            parts.append(f"<b>{k}:</b> {v}")
        return Paragraph(" ".join(parts), kv_s)

    def _tbl(headers, rows, col_widths=None):
        data = [[str(c) for c in headers]] + [[str(c) for c in r] for r in rows]
        cw   = col_widths or [W / len(headers)] * len(headers)
        t = Table(data, colWidths=cw, repeatRows=1)
        t.setStyle(TableStyle([
            ("BACKGROUND",    (0, 0), (-1, 0),  BLUE),
            ("TEXTCOLOR",     (0, 0), (-1, 0),  colors.white),
            ("FONTNAME",      (0, 0), (-1, 0),  "Helvetica-Bold"),
            ("FONTNAME",      (0, 1), (-1, -1), "Helvetica"),
            ("FONTSIZE",      (0, 0), (-1, -1), 7.5),
            ("ROWBACKGROUNDS",(0, 1), (-1, -1), [colors.white, colors.HexColor("#eef4fb")]),
            ("GRID",          (0, 0), (-1, -1), 0.3, colors.HexColor("#c8ddf0")),
            ("VALIGN",        (0, 0), (-1, -1), "MIDDLE"),
            ("TOPPADDING",    (0, 0), (-1, -1), 3),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
        ]))
        return t

    def _img2col(key1, key2):
        if key1 not in charts and key2 not in charts:
            return None
        cells = []
        for key in (key1, key2):
            if key in charts:
                cells.append([Image(charts[key], width=W*0.47, height=W*0.32)])
            else:
                cells.append([""])
        t = Table([cells], colWidths=[W*0.5, W*0.5])
        t.setStyle(TableStyle([
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
            ("ALIGN",  (0, 0), (-1, -1), "CENTER"),
        ]))
        return t

    story = []

    # ── PAGE 1: Cover ─────────────────────────────────────────────────────────
    story.append(Spacer(1, 1.0*inch))
    if os.path.exists(LOGO_PATH):
        story.append(Image(LOGO_PATH, width=1.8*inch, height=1.8*inch, kind="proportional"))
        story.append(Spacer(1, 0.12*inch))
    story.append(Paragraph("Sewa Sanskriti App", _s("pH", fontSize=18, leading=22,
                 fontName="Helvetica-Bold", textColor=BLUE, alignment=TA_CENTER)))
    story.append(Spacer(1, 0.08*inch))
    story.append(Paragraph("Case Study Report", title_s))
    story.append(Paragraph(scope_line, sub_s))
    story.append(Spacer(1, 0.06*inch))
    all_mat_pdf = ", ".join(ml["top_matrusansthan"].keys())
    story.append(Paragraph(f"Matrusansthans: {all_mat_pdf}", _s("pMatPDF", fontSize=8.5,
                 leading=12, alignment=TA_CENTER, fontName="Helvetica-Oblique")))
    story.append(Spacer(1, 0.08*inch))
    story.append(Paragraph(
        f"Rashtriya Sewa Bharti (RSB)  |  Scope: {ml['scope_level']} — {scope_line}  |  "
        f"Date: {date.today().strftime('%B %d, %Y')}",
        meta_s,
    ))
    story.append(PageBreak())

    # ── PAGE 2: Overview + qualitative + 2 charts ─────────────────────────────
    story += _sec_elem("Overview")
    story.append(_kvrow_elem([
        ("Prakalps", f"{ml['total_records']:,}"),
        ("Verified",  f"{ml['verified']:,} ({ml['verification_rate']}%)"),
        ("Districts", ml["districts_covered"]),
        ("Villages",  ml["villages_covered_unique"]),
    ]))
    story.append(_kvrow_elem([
        ("Cumulative Beneficiaries", f"{ml['ben_total_ever']:,}"),
        ("Incharge/Workers", f"{ml.get('incharge_count', ml['vol_total']):,}"),
        ("Gram Samiti", f"{ml['gram_male'] + ml['gram_female']:,}"),
        ("Families",   f"{ml['families_covered']:,}"),
    ]))
    mat_names_pdf = ", ".join(f"{k} ({v})" for k, v in ml["top_matrusansthan"].items())
    story.append(_kvrow_elem([
        ("Dominant Vertical",  ml["top_vertical"]),
        ("Female Share",       f"{ml['female_pct']}%"),
        ("Matrusansthans Involved", f"{ml['matrusansthan_count']} — {mat_names_pdf}"),
    ]))

    for label in ["Reasons for Opening", "Reported Impact", "Achievements"]:
        data = ml["text_analysis"].get(label)
        if not data:
            continue
        story += _sec_elem(label)
        for s in (data.get("samples") or [])[:2]:
            story.append(Paragraph(f"- {s}", italic))
        if data["top_themes"]:
            story.append(_kvrow_elem([("Key themes",
                ", ".join(kw for kw, _ in data["top_themes"]))]))

    img2 = _img2col("gender_chart", "vertical_chart")
    if img2:
        story.append(Spacer(1, 4))
        story.append(img2)

    # ── PAGE 3: HR + Reach + Verticals + Operational ──────────────────────────
    story += _sec_elem("Human Resources")
    hr_rows = []
    if ml["vol_total"] > 0:
        hr_rows.append(("Volunteers", ml["vol_male"], ml["vol_female"], ml["vol_total"]))
    if ml["sal_total"] > 0:
        hr_rows.append(("Salaried Staff", ml["sal_male"], ml["sal_female"], ml["sal_total"]))
    gram_total = ml["gram_male"] + ml["gram_female"]
    if gram_total > 0:
        hr_rows.append(("Gram Samiti", ml["gram_male"], ml["gram_female"], gram_total))
    if ml.get("incharge_count", 0) > 0:
        hr_rows.append(("Incharge / Coordinators", ml["incharge_count"], "-", ml["incharge_count"]))
    if hr_rows:
        story.append(_tbl(["Category", "Male", "Female", "Total"], hr_rows))
    else:
        story.append(_kvrow_elem([("Note", "Worker breakdown not available in this dataset")]))

    story += _sec_elem("Beneficiary Reach")
    story.append(_tbl(
        ["Category", "Current", "Cumulative"],
        [
            ("Male",   f"{ml['ben_male_now']:,}",   f"{ml['ben_male_ever']:,}"),
            ("Female", f"{ml['ben_female_now']:,}", f"{ml['ben_female_ever']:,}"),
            ("Boys",   f"{ml['ben_boys_now']:,}",   f"{ml['ben_boys_ever']:,}"),
            ("Girls",  f"{ml['ben_girls_now']:,}",  f"{ml['ben_girls_ever']:,}"),
            ("TOTAL",  f"{ml['ben_total_now']:,}",  f"{ml['ben_total_ever']:,}"),
        ],
    ))

    story += _sec_elem("Service Verticals")
    story.append(_tbl(
        ["Vertical", "Prakalps"],
        [(k, str(v)) for k, v in list(ml["verticals"].items())[:6]],
        col_widths=[W*0.75, W*0.25],
    ))

    # ── PAGE 4: Impact + Segments + Summary ───────────────────────────────────
    story += _sec_elem("Top Impact Prakalps")
    story.append(_kvrow_elem([
        ("Avg Score",   ml["avg_impact_score"]),
        ("High-Impact", ml["high_impact_count"]),
        ("Outliers",    ml["outlier_count"]),
    ]))
    if ml["top_projects"]:
        story.append(_tbl(
            ["Prakalp Name", "Vertical", "Total Beneficiaries", "Score"],
            [(p["prakalp_name"][:28], p["vertical_name"],
              f"{int(p['total_beneficiaries_till_date']):,}",
              str(round(p["impact_score"], 2)))
             for p in ml["top_projects"][:5]],
            col_widths=[W*0.38, W*0.28, W*0.20, W*0.14],
        ))

    if ml["segments"]:
        story += _sec_elem("Prakalp Segments")
        story.append(_tbl(
            ["Segment", "Prakalps", "Avg Beneficiaries", "Top Vertical"],
            [(s["label"], str(s["count"]), str(round(s["avg_ben"], 0)), s["top_vertical"])
             for s in ml["segments"]],
            col_widths=[W*0.22, W*0.15, W*0.18, W*0.45],
        ))

    story += _sec_elem("Recommendations")
    for i, rec in enumerate(_get_recommendations(ml)[:5], 1):
        story.append(Paragraph(f"{i}.  {rec}", normal))

    story += _sec_elem("Summary")
    n = ml["total_records"]
    centre_word_pdf = "centres" if n > 1 else "centre"
    mat_names_pdf2  = ", ".join(ml["top_matrusansthan"].keys())
    verticals_pdf   = ", ".join(list(ml["verticals"].keys())[:4])
    summary_pdf = (
        f"This {ml['scope_level'].lower()}-level report covers {n:,} prakalps across "
        f"{ml['districts_covered']} district(s) and {ml['villages_covered_unique']} village(s) in {scope_line}. "
        f"These {centre_word_pdf} are run under {ml['matrusansthan_count']} matrusansthans — {mat_names_pdf2}. "
        f"Cumulatively, {ml['ben_total_ever']:,} beneficiaries have been served "
        f"({ml['ben_total_now']:,} currently active), with a female share of {ml['female_pct']}%. "
        f"The dominant vertical is {ml['top_vertical']}, also covering {verticals_pdf}. "
        f"{ml['high_impact_count']} prakalps were identified as high-impact. "
        f"{ml['verified']:,} of {n:,} prakalps ({ml['verification_rate']}%) are verified."
    )
    for lbl in ["Reported Impact", "Achievements", "Reasons for Opening"]:
        data = ml["text_analysis"].get(lbl)
        if data and data.get("samples"):
            summary_pdf += f" {lbl}: {data['samples'][0]}"
            break
    story.append(Paragraph(summary_pdf, normal))

    # ── PAGE 5: Photos (2×2 grid, max 4) ──────────────────────────────────────
    if photos:
        story.append(PageBreak())
        story += _sec_elem("Photo Gallery")
        story.append(Paragraph("Field photographs from project sites.", normal))
        for row_start in range(0, min(len(photos), 4), 2):
            pair = photos[row_start:row_start + 2]
            cells = []
            for _, cap, tmp_path in pair:
                try:
                    cells.append([
                        Image(tmp_path, width=W*0.46, height=W*0.32),
                        Paragraph(_clean(cap), cap_s),
                    ])
                except Exception:
                    cells.append([""])
            while len(cells) < 2:
                cells.append([""])
            pt = Table([cells], colWidths=[W*0.5, W*0.5])
            pt.setStyle(TableStyle([
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("ALIGN",  (0, 0), (-1, -1), "CENTER"),
            ]))
            story.append(pt)
            story.append(Spacer(1, 6))

    doc.build(story)
    return tmp.name

def _get_recommendations(ml: dict) -> list:
    recs = []
    vmap = {
        "Education": "Scale tuition centres and digital literacy; target school dropout age groups.",
        "Health": "Expand Arogya Rakshak coverage; run quarterly health camps.",
        "Self-Reliance": "Introduce vocational training and micro-finance linkages.",
        "Social/Cultural": "Leverage cultural events for youth engagement and community cohesion.",
    }
    for key, rec in vmap.items():
        if any(key.lower() in v.lower() for v in ml["verticals"]):
            recs.append(f"[{key}] {rec}")
    if ml["high_impact_count"] > 0:
        recs.append(
            f"Document and replicate the {ml['high_impact_count']} high-impact project models "
            f"across similar villages."
        )
    if ml["outlier_count"] > 0:
        recs.append(f"Audit {ml['outlier_count']} outlier projects for data accuracy or special needs.")
    if ml["female_pct"] < 40:
        recs.append("Increase female beneficiary and volunteer participation — current share is below 40%.")
    if ml["verification_rate"] < 80:
        recs.append(
            f"Improve data verification rate (currently {ml['verification_rate']}%) "
            f"through field follow-ups."
        )
    if ml["similar_units"]:
        sv = ml["similar_units"][0]
        col = ml["similar_col"]
        recs.append(
            f"Benchmark against {sv[col]} (similarity {round(sv['similarity'], 2)}) "
            f"to adopt proven practices."
        )
    recs.append("Establish quarterly impact reviews to track score trends and beneficiary growth.")
    return recs
